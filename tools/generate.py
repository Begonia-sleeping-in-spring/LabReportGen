import os
from docx import Document

BASE = "experiments"
REPORTS = "reports"


def latest():
    exps = sorted(
        [x for x in os.listdir(BASE)
         if x.startswith("exp_")]
    )

    if not exps:
        raise Exception("No experiment found")

    return os.path.join(BASE, exps[-1])


def read_file(path):
    if os.path.exists(path):
        with open(path, "r", errors="ignore") as f:
            return f.read()
    return ""


exp = latest()
exp_name = os.path.basename(exp)

metadata = read_file(f"{exp}/metadata.conf")
commands = read_file(f"{exp}/commands.log")
events = read_file(f"{exp}/events.log")
files = read_file(f"{exp}/files.log")

os.makedirs(REPORTS, exist_ok=True)

# -------------------------
# Markdown
# -------------------------

md_out = f"{REPORTS}/{exp_name}.md"

with open(md_out, "w") as f:

    f.write(f"# Experiment Report\n\n")
    f.write(f"Experiment: {exp_name}\n\n")

    f.write("## Metadata\n\n")
    f.write("```\n")
    f.write(metadata)
    f.write("\n```\n\n")

    f.write("## Commands\n\n")
    f.write("```\n")
    f.write(commands)
    f.write("\n```\n\n")

    f.write("## Events\n\n")
    f.write("```\n")
    f.write(events)
    f.write("\n```\n\n")

    f.write("## File Snapshots\n\n")
    f.write("```\n")
    f.write(files)
    f.write("\n```\n\n")

# -------------------------
# DOCX
# -------------------------

doc = Document()

doc.add_heading("LabReport v1.7", 0)

doc.add_heading("Experiment", level=1)
doc.add_paragraph(exp_name)

doc.add_heading("Metadata", level=1)
doc.add_paragraph(metadata)

doc.add_heading("Commands", level=1)
doc.add_paragraph(commands)

doc.add_heading("Events", level=1)
doc.add_paragraph(events)

doc.add_heading("File Snapshots", level=1)
doc.add_paragraph(files)

img_dir = f"{exp}/screenshots"

if os.path.exists(img_dir):

    doc.add_heading("Screenshots", level=1)

    for img in sorted(os.listdir(img_dir)):
        doc.add_paragraph(img)

docx_out = f"{REPORTS}/{exp_name}.docx"

doc.save(docx_out)

print("Markdown:", md_out)
print("DOCX:", docx_out)
